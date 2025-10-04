# ***************************************************************************
# *   Copyright (c) 2025 David Carter <dcarter@davidcarter.ca>              *
# *                                                                         *
# *   This program is free software; you can redistribute it and/or modify  *
# *   it under the terms of the GNU Lesser General Public License (LGPL)    *
# *   as published by the Free Software Foundation; either version 2 of     *
# *   the License, or (at your option) any later version.                   *
# *   for detail see the LICENCE text file.                                 *
# *                                                                         *
# *   This program is distributed in the hope that it will be useful,       *
# *   but WITHOUT ANY WARRANTY; without even the implied warranty of        *
# *   MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the         *
# *   GNU Library General Public License for more details.                  *
# *                                                                         *
# *   You should have received a copy of the GNU Library General Public     *
# *   License along with this program; if not, write to the Free Software   *
# *   Foundation, Inc., 59 Temple Place, Suite 330, Boston, MA  02111-1307  *
# *   USA                                                                   *
# *                                                                         *
# ***************************************************************************
"""Class for CFD Reports"""

__author__ = "David Carter"
__url__ = "https://www.davesrocketshop.com"

import FreeCAD
import os
import csv
import math

import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvas

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from CfdOF import CfdTools
from CfdOF.Mesh.CfdMesh import MESHER_DESCRIPTIONS, MESHERS
from CfdOF.Mesh.CfdMeshRefinement import EXTRUSION_NAMES, EXTRUSION_TYPES
from CfdOF.Solve.CfdFluidBoundary import BOUNDARY_NAMES, BOUNDARY_TYPES, SUBNAMES, SUBTYPES, TURBULENT_INLET_SPEC
from CfdOF.Solve.TaskPanelCfdFluidProperties import ALL_FIELDS

from Rocket.Utilities import _msg
from Rocket.cfd.FeatureCFDRocket import applyTranslations, calcFrontalArea

class CFDReport:

    def __init__(self, analysis):
        # Import here to prevent circular imports
        from Rocket.cfd.CFDUtil import caliber, finThickness
        from Rocket.cfd.FeatureCFDRocket import calcFrontalArea

        self._analysis = analysis
        self._path = os.path.join(CfdTools.getOutputPath(self._analysis), 'CFD Report.docx')

        self._frontalArea = calcFrontalArea(self._analysis.Shape)
        self._diameter = caliber(self._analysis.Rocket)
        self._thickness = finThickness(self._analysis.Rocket)
        box = self._analysis.Shape.BoundBox
        self._length = box.XLength
        self._x0 = self._length / 2.0
        self._rotation = self._analysis.CFDRocket.AngleOfRotation.Value

        self._totalRunTime = ''

        self._runStatus = {}
        self._forces = {}
        self._moments = {}
        self._coefficients = {}

    def getPath(self):
        return self._path

    def collectStats(self):
        path = os.path.join(CfdTools.getOutputPath(self._analysis), "RunStatus.dat")
        self.collectStatusInformation(path)

        for angle in self._analysis.AOAList:
            try:
                if self._runStatus[str(angle)][2] == "Success":
                    dirname = "case_aoa_{}".format(angle)

                    path = os.path.join(CfdTools.getOutputPath(self._analysis), dirname,
                                        "postProcessing", "ForceReportingFunction", "0", "force.dat")
                    self.collectForceInformation(path, angle)

                    path = os.path.join(CfdTools.getOutputPath(self._analysis), dirname,
                                        "postProcessing", "ForceReportingFunction", "0", "moment.dat")
                    self.collectMomentInformation(path, angle)

                    path = os.path.join(CfdTools.getOutputPath(self._analysis), dirname,
                                        "postProcessing", "ForceCoefficientReportingFunction", "0", "coefficient.dat")
                    self.collectCoefficientInformation(path, angle)
            except KeyError:
                pass

    def collectStatusInformation(self, path):
        with open(path, "r") as csvfile:
            csvreader = csv.reader(csvfile, delimiter='\t', skipinitialspace=True)
            for row in csvreader:
                if len(row) > 2:
                    self._runStatus[str(row[0])] = row
                elif len(row) == 2:
                    if row[0] == "Total":
                        self._totalRunTime = row[1]

    def collectForceInformation(self, path, angle):
        self._forces[str(angle)] = []
        with open(path, "r") as csvfile:
            csvreader = csv.reader(csvfile, delimiter=' ', skipinitialspace=True)
            for row in csvreader:
                if len(row) == 10 and not row[0].startswith("#"):
                    self._forces[str(angle)].append(row)

    def collectMomentInformation(self, path, angle):
        self._moments[str(angle)] = []
        with open(path, "r") as csvfile:
            csvreader = csv.reader(csvfile, delimiter=' ', skipinitialspace=True)
            for row in csvreader:
                # Read all and save the last row
                if len(row) == 10 and not row[0].startswith("#"):
                    self._moments[str(angle)].append(row)

    def collectCoefficientInformation(self, path, angle):
        self._coefficients[str(angle)] = []
        with open(path, "r") as csvfile:
            csvreader = csv.reader(csvfile, delimiter='\t', skipinitialspace=True)
            for row in csvreader:
                # Read all and save the last row
                if len(row) == 13 and not row[0].startswith("#"):
                    self._coefficients[str(angle)].append(row)

    def generate(self):
        self.collectStats()

        self._document = Document()
        self.addStyles()

        self.generateIntro()
        self.generateRuntime()
        self.generateCP()
        self.generateCD()
        self.generateRunData()
        self._document.save(self._path)

    def addStyles(self):
        """ Add some LibreOffice styles """
        styles = self._document.styles
        styles.add_style('Block Quotation', WD_STYLE_TYPE.PARAGRAPH, builtin = True)
        styles.add_style('Simple Grid Columns', WD_STYLE_TYPE.TABLE, builtin = True)

    def generateIntro(self):
        self._document.add_heading('CFD Report', 0)

        p = self._document.add_paragraph("This is an automated report generated by the Rocket Workbench. "\
                                         "It was created by running a CFD analysis on the document '{}'."\
                                            .format(FreeCAD.ActiveDocument.Label))

        self._document.add_paragraph("Details about the study, including the rocket")

        table = self._document.add_table(rows=4, cols=2, style='Table Grid')
        hdr_cells = table.columns[0].cells
        hdr_cells[0].text = 'Length'
        hdr_cells[1].text = 'Diameter'
        hdr_cells[2].text = 'Minimum Fin Thickness'
        hdr_cells[3].text = 'Frontal Area'
        data_cells = table.columns[1].cells
        data_cells[0].text = FreeCAD.Units.Quantity("{} mm".format(self._length)).UserString
        data_cells[1].text = FreeCAD.Units.Quantity("{} mm".format(self._diameter)).UserString
        data_cells[2].text = FreeCAD.Units.Quantity("{} mm".format(self._thickness)).UserString
        data_cells[3].text = FreeCAD.Units.Quantity("{} mm^2".format(self._frontalArea)).UserString

        self._document.add_heading('Study Parameters', level=1)

        self._document.add_paragraph("This study varied the following parameters:")
        table = self._document.add_table(rows=1, cols=1, style='Table Grid')
        table.autofit = True
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Angle of Attack'
        for angle in self._analysis.AOAList:
            row_cells = table.add_row().cells
            row_cells[0].text = str(angle)

        non_zero = 0
        for angle in self._analysis.AOAList:
            if angle != 0:
                non_zero += 1
        if non_zero < 2:
            p = self._document.add_paragraph()
            p.add_run("NOTE: It is not possible to calculate the center of pressure at an "\
                                        "angle of attack of 0. It must be calculated at multiple points close "\
                                        "to 0 and inferred using l'Hôpital's rule.").bold = True
            p = self._document.add_paragraph()
            p.add_run('This study is unable to determine Center of Pressure at a 0 angle of attack.').italic = True

        self._document.add_heading('Configuration', level=1)

        self.physicsModelConfig()
        self.fluidPropertiesConfig()
        self.initializeFlowConfig()
        self.solverConfig()
        self.mesherConfig()
        self.boundaryConditionsConfig()

        self._document.add_page_break()

    def physicsModelConfig(self):
        self._document.add_heading('Physics Model', level=2)

        physicsModel = CfdTools.getPhysicsModel(self._analysis)
        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True

        row_cells = table.rows[0].cells
        row_cells[0].text = "Time"
        row_cells[1].text = physicsModel.Time

        row_cells = table.add_row().cells
        row_cells[0].text = "Flow"
        row_cells[1].text = physicsModel.Phase

        row_cells = table.add_row().cells
        row_cells[0].text = "Isothermal"
        row_cells[1].text = self.configState(physicsModel.Flow == "Isothermal")

        row_cells = table.add_row().cells
        row_cells[0].text = "High Mach number"
        row_cells[1].text = self.configState(physicsModel.Flow == "HighMachCompressible")

        row_cells = table.add_row().cells
        row_cells[0].text = "Viscous"
        row_cells[1].text = self.configState(physicsModel.Turbulence == "Inviscid")

        row_cells = table.add_row().cells
        row_cells[0].text = "Rotating frame (SRF)"
        row_cells[1].text = self.configState(physicsModel.SRFModelEnabled)

        if physicsModel.Turbulence != "Inviscid":
            row_cells = table.add_row().cells
            row_cells[0].text = "Turbulence"
            row_cells[1].text = physicsModel.Turbulence
            if physicsModel.Turbulence != "Laminar":
                row_cells = table.add_row().cells
                row_cells[0].text = "Turbulence Model"
                row_cells[1].text = physicsModel.TurbulenceModel

        if physicsModel.Flow != "Isothermal" or physicsModel.Phase == "FreeSurface":
            row_cells = table.add_row().cells
            row_cells[0].text = "Gravity (x, y, z)"
            row_cells[1].text = "({}, {}, {})".format(
                physicsModel.gx.UserString,
                physicsModel.gy.UserString,
                physicsModel.gz.UserString,
            )

        if physicsModel.SRFModelEnabled and physicsModel.Phase != "FreeSurface":
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].text = "Moving Reference Frame (SRF)"

            row_cells = table.add_row().cells
            row_cells[0].text = "RPM"
            row_cells[1].text = physicsModel.SRFModelRPM.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Center of Rotation"
            row_cells[1].text = "({} mm, {} mm, {} mm)".format(
                physicsModel.SRFModelCoR.x,
                physicsModel.SRFModelCoR.y,
                physicsModel.SRFModelCoR.z,
            )

            row_cells = table.add_row().cells
            row_cells[0].text = "Rotational Axis"
            row_cells[1].text = "({}, {}, {})".format(
                physicsModel.SRFModelAxis.x,
                physicsModel.SRFModelAxis.y,
                physicsModel.SRFModelAxis.z,
            )

    def fluidPropertiesConfig(self):
        self._document.add_heading('Fluid Properties', level=2)

        fluidModel = CfdTools.getMaterials(self._analysis)
        if len(fluidModel) == 1:
            fluidModel = fluidModel[0]
        else:
            self._document.add_paragraph("Not defined.")
            return

        material_type =  fluidModel.Material["Type"]

        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True

        row_cells = table.rows[0].cells
        row_cells[0].text = "Material Name"
        row_cells[1].text = fluidModel.Label

        row_cells = table.add_row().cells
        row_cells[0].text = "Type"
        row_cells[1].text = material_type

        fields = ALL_FIELDS[material_type]
        for name in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = fluidModel.Material.get(name, '0')

    def initializeFlowConfig(self):
        self._document.add_heading('Initialize Flow Field', level=2)

        initModel = CfdTools.getInitialConditions(self._analysis)
        physicsModel = CfdTools.getPhysicsModel(self._analysis)

        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True

        row_cells = table.rows[0].cells
        row_cells[0].text = "Velocity"
        if initModel.PotentialFlow:
            row_cells[1].text = "Potential flow"
        elif initModel.UseInletUValues:
            row_cells[1].text = "Use values from boundary '{}'".format(initModel.BoundaryU.Label)
        else:
            row_cells[1].text = "({}, {}, {})".format(
                initModel.Ux.UserString,
                initModel.Uy.UserString,
                initModel.Uz.UserString,
            )

        row_cells = table.add_row().cells
        row_cells[0].text = "Pressure"
        if initModel.PotentialFlowP:
            row_cells[1].text = "Potential flow"
        elif initModel.UseOutletPValue:
            row_cells[1].text = "Use values from boundary '{}'".format(initModel.BoundaryP.Label)
        else:
            row_cells[1].text = initModel.Pressure.UserString

        if physicsModel.Phase != 'Single':
            row_cells = table.add_row().cells
            row_cells[0].text = "Volume Fractions"

            if len(initModel.VolumeFractions) == 0:
                _msg("No fractions to report")
                row_cells[1].text = "None"
            else:
                row_cells[0].merge(row_cells[1])

            for name, value in initModel.VolumeFractions.items():
                _msg("key {}, value {}".format(name, value))
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = value

        if physicsModel.Flow != "Isothermal":
            row_cells = table.add_row().cells
            if initModel.UseInletTemperatureValue:
                row_cells[0].text = "Thermal"
                row_cells[1].text = "Use values from boundary '{}'".format(initModel.BoundaryT.Label)
            else:
                row_cells[0].text = "Temperature"
                row_cells[1].text = initModel.Temperature.UserString

        row_cells = table.add_row().cells
        row_cells[0].text = "Turbulence"
        if initModel.UseInletTurbulenceValues:
            row_cells[1].text = "Use values from boundary '{}'".format(initModel.BoundaryTurb.Label)
        else:
            row_cells[0].merge(row_cells[1])

            row_cells = table.add_row().cells
            row_cells[0].text = "Model" # This is a duplicate of the Physics model value
            row_cells[1].text = physicsModel.TurbulenceModel

            if physicsModel.TurbulenceModel in ['kOmegaSST', 'kOmegaSSTLM', 'kOmegaSSTDES', 'kEpsilon', 'kEqn']:
                row_cells = table.add_row().cells
                row_cells[0].text = "k"
                row_cells[1].text = initModel.k.UserString

            if physicsModel.TurbulenceModel in ['kEpsilon']:
                row_cells = table.add_row().cells
                row_cells[0].text = "ε"
                row_cells[1].text = initModel.epsilon.UserString

            if physicsModel.TurbulenceModel in ['kOmegaSST', 'kOmegaSSTLM', 'kOmegaSSTDES']:
                row_cells = table.add_row().cells
                row_cells[0].text = "ω"
                row_cells[1].text = initModel.omega.UserString

            if physicsModel.TurbulenceModel in ['SpalartAllmaras']:
                row_cells = table.add_row().cells
                row_cells[0].text = "ṽ"
                row_cells[1].text = initModel.nuTilda.UserString

            if physicsModel.TurbulenceModel in ['kOmegaSSTLM']:
                row_cells = table.add_row().cells
                row_cells[0].text = "γ"
                row_cells[1].text = initModel.gammaInt.UserString

                row_cells = table.add_row().cells
                row_cells[0].text = "Reθ"
                row_cells[1].text = initModel.ReThetat.UserString

            if physicsModel.TurbulenceModel in ['Smagorinsky', 'WALE', 'kEqn']:
                row_cells = table.add_row().cells
                row_cells[0].text = "v"
                row_cells[1].text = initModel.nut.UserString

    def solverConfig(self):
        self._document.add_heading('Solver', level=2)

        solver = CfdTools.getSolver(self._analysis)
        physicsModel = CfdTools.getPhysicsModel(self._analysis)

        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True

        row_cells = table.rows[0].cells
        row_cells[0].text = "Convergence tolerance"
        row_cells[1].text = str(solver.ConvergenceTol)

        row_cells = table.add_row().cells
        row_cells[0].text = "Max iterations"
        row_cells[1].text = str(solver.MaxIterations)

        row_cells = table.add_row().cells
        row_cells[0].text = "Iteration write interval"
        row_cells[1].text = str(solver.SteadyWriteInterval)

        if solver.Parallel:
            row_cells = table.add_row().cells
            row_cells[0].text = "Parallel cores"
            row_cells[1].text = str(solver.ParallelCores)

        row_cells = table.add_row().cells
        row_cells[0].text = "Starting from"
        row_cells[1].text = solver.StartFrom

        if physicsModel.Time == 'Transient':
            row_cells = table.add_row().cells
            row_cells[0].text = "End time"
            row_cells[1].text = solver.EndTime.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Time step"
            row_cells[1].text = solver.EndTime.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Max CFL number"
            row_cells[1].text = float(solver.MaxCFLNumber)

            row_cells = table.add_row().cells
            row_cells[0].text = "Max free-surface CFL number"
            row_cells[1].text = float(solver.MaxInterfaceCFLNumber)

            row_cells = table.add_row().cells
            row_cells[0].text = "Transient write interval"
            row_cells[1].text = solver.TransientWriteInterval.UserString

    def mesherConfig(self):
        self._document.add_heading('Mesher', level=2)

        mesher = CfdTools.getMeshObject(self._analysis)

        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True

        row_cells = table.rows[0].cells
        row_cells[0].text = "Mesh utility"
        row_cells[1].text = MESHER_DESCRIPTIONS[MESHERS.index(mesher.MeshUtility)]

        row_cells = table.add_row().cells
        row_cells[0].text = "Base element size"
        row_cells[1].text = mesher.CharacteristicLengthMax.UserString

        if mesher.MeshUtility == 'snappyHexMesh':
            row_cells = table.add_row().cells
            row_cells[0].text = "Cells between levels"
            row_cells[1].text = str(mesher.CellsBetweenLevels)

            row_cells = table.add_row().cells
            row_cells[0].text = "Relative edge refinement"
            row_cells[1].text = str(mesher.EdgeRefinement)

            row_cells = table.add_row().cells
            row_cells[0].text = "Edge detection"
            if mesher.ImplicitEdgeDetection:
                row_cells[1].text = "Implicit"
            else:
                row_cells[1].text = "Explicit"

        row_cells = table.add_row().cells
        row_cells[0].text = "STL relative linear deflection"
        row_cells[1].text = str(mesher.STLRelativeLinearDeflection)

        row_cells = table.add_row().cells
        row_cells[0].text = "STL angulat mesh density"
        row_cells[1].text = str(mesher.STLAngularMeshDensity)

        row_cells = table.add_row().cells
        row_cells[0].text = "Convert to dual mesh"
        row_cells[1].text = self.configState(mesher.ConvertToDualMesh)

        self._document.add_heading('Refinements', level=3)
        for refinement in mesher.Group:
            self.refinementConfig(refinement)

    def refinementConfig(self, refinement):

        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True

        row_cells = table.rows[0].cells
        row_cells[0].text = "Name"
        row_cells[1].text = refinement.Label

        row_cells = table.add_row().cells
        row_cells[0].text = "Refinement type"
        if refinement.Internal:
            row_cells[1].text = "Volume"

            row_cells = table.add_row().cells
            row_cells[0].text = "Relative element size"
            row_cells[1].text = str(refinement.RelativeLength)
        elif refinement.Extrusion:
            row_cells[1].text = "Extrusion"

            row_cells = table.add_row().cells
            row_cells[0].text = "Extrusion type"
            row_cells[1].text = EXTRUSION_NAMES[EXTRUSION_TYPES.index(refinement.ExtrusionType)]

            if refinement.ExtrusionType == "2DPlanar":
                row_cells = table.add_row().cells
                row_cells[0].text = "Thickness"
                row_cells[1].text = refinement.ExtrusionThickness.UserString

            elif refinement.ExtrusionType == "2DWedge":
                row_cells = table.add_row().cells
                row_cells[0].text = "Axis point"
                row_cells[1].text = "({} mm, {} mm, {} mm)".format(
                    refinement.ExtrusionAxisPoint.x,
                    refinement.ExtrusionAxisPoint.y,
                    refinement.ExtrusionAxisPoint.z,
                )

                row_cells = table.add_row().cells
                row_cells[0].text = "Axis direction"
                row_cells[1].text = "({}, {}, {})".format(
                    refinement.ExtrusionAxisDirection.x,
                    refinement.ExtrusionAxisDirection.y,
                    refinement.ExtrusionAxisDirection.z,
                )

            elif refinement.ExtrusionType == "PatchNormal":
                row_cells = table.add_row().cells
                row_cells[0].text = "Keep existing mesh"
                row_cells[1].text = self.configState(refinement.KeepExistingMesh)

                row_cells = table.add_row().cells
                row_cells[0].text = "Thickness"
                row_cells[1].text = refinement.ExtrusionThickness.UserString

                row_cells = table.add_row().cells
                row_cells[0].text = "Number of layers"
                row_cells[1].text = str(refinement.ExtrusionLayers)

                row_cells = table.add_row().cells
                row_cells[0].text = "Expansion ratio"
                row_cells[1].text = str(refinement.ExtrusionRatio)

            elif refinement.ExtrusionType == "Rotational":
                row_cells = table.add_row().cells
                row_cells[0].text = "Keep existing mesh"
                row_cells[1].text = self.configState(refinement.KeepExistingMesh)

                row_cells = table.add_row().cells
                row_cells[0].text = "Angle"
                row_cells[1].text = refinement.ExtrusionAngle.UserString

                row_cells = table.add_row().cells
                row_cells[0].text = "Number of layers"
                row_cells[1].text = str(refinement.ExtrusionLayers)

                row_cells = table.add_row().cells
                row_cells[0].text = "Expansion ratio"
                row_cells[1].text = str(refinement.ExtrusionRatio)

                row_cells = table.add_row().cells
                row_cells[0].text = "Axis point"
                row_cells[1].text = "({} mm, {} mm, {} mm)".format(
                    refinement.ExtrusionAxisPoint.x,
                    refinement.ExtrusionAxisPoint.y,
                    refinement.ExtrusionAxisPoint.z,
                )

                row_cells = table.add_row().cells
                row_cells[0].text = "Axis direction"
                row_cells[1].text = "({}, {}, {})".format(
                    refinement.ExtrusionAxisDirection.x,
                    refinement.ExtrusionAxisDirection.y,
                    refinement.ExtrusionAxisDirection.z,
                )

        else:
            row_cells[1].text = "Surface"

            row_cells = table.add_row().cells
            row_cells[0].text = "Relative element size"
            row_cells[1].text = str(refinement.RelativeLength)

            row_cells = table.add_row().cells
            row_cells[0].text = "Refinement thickness"
            row_cells[1].text = refinement.RefinementThickness.UserString

            if refinement.NumberLayers > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = "Boundary layers"
                row_cells[1].text = str(refinement.NumberLayers)

                row_cells = table.add_row().cells
                row_cells[0].text = "Expansion ratio"
                row_cells[1].text = str(refinement.ExpansionRatio)

                row_cells = table.add_row().cells
                row_cells[0].text = "Max first cell height"
                row_cells[1].text = refinement.FirstLayerHeight.UserString

        # for shape in refinement.ShapeRefs:
        #     for subShape in shape:
        #         _msg(str(dir(subShape)))


        self._document.add_paragraph() # Acts as a spacer between tables

    def boundaryConditionsConfig(self):
        self._document.add_heading('Boundary Conditions', level=2)

        boundaries = CfdTools.getCfdBoundaryGroup(self._analysis)
        for bc in boundaries:
            self.bcConfig(bc)

    def bcConfig(self, bc):

        table = self._document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = True
        # from CfdOF.Solve.CfdFluidBoundary import BOUNDARY_NAMES, BOUNDARY_TYPES, SUBNAMES, SUBTYPES

        row_cells = table.rows[0].cells
        row_cells[0].text = "Name"
        row_cells[1].text = bc.Label

        typeIndex = BOUNDARY_TYPES.index(bc.BoundaryType)
        row_cells = table.add_row().cells
        row_cells[0].text = "Type"
        row_cells[1].text = BOUNDARY_NAMES[typeIndex]

        row_cells = table.add_row().cells
        row_cells[0].text = "Sub type"
        row_cells[1].text = SUBNAMES[typeIndex][SUBTYPES[typeIndex].index(bc.BoundarySubType)]

        row_cells = table.add_row().cells
        row_cells[0].text = "Default boundary"
        row_cells[1].text = self.configState(bc.DefaultBoundary)

        if typeIndex == 0:
            self.bcWallConfig(table, typeIndex, bc)
        elif typeIndex == 1:
            self.bcInletConfig(table, typeIndex, bc)
        elif typeIndex == 2:
            self.bcOutletConfig(table, typeIndex, bc)
        elif typeIndex == 3:
            self.bcOpenConfig(table, typeIndex, bc)
        elif typeIndex == 4:
            self.bcConstraintConfig(table, typeIndex, bc)
        elif typeIndex == 5:
            self.bcBaffleConfig(table, typeIndex, bc)

        self._document.add_paragraph() # Acts as a spacer between tables

    def bcWallConfig(self, table, typeIndex, bc):
        if bc.BoundarySubType == "partialSlipWall":
            row_cells = table.add_row().cells
            row_cells[0].text = "Slip ratio"
            row_cells[1].text = bc.SlipRatio.UserString

        if bc.BoundarySubType == "translatingWall" or bc.BoundarySubType == "roughWall":
            self.bcVelocityConfig(table, bc)

        if bc.BoundarySubType == "roughWall":
            row_cells = table.add_row().cells
            row_cells[0].text = "Roughness height (Ks)"
            row_cells[1].text = bc.RoughnessHeight.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Roughness constant (Cs)"
            row_cells[1].text = bc.RoughnessConstant.UserString

    def bcInletConfig(self, table, typeIndex, bc):
        if bc.BoundarySubType == "uniformVelocityInlet" or bc.BoundarySubType == "staticPressureInlet":
            self.bcVelocityConfig(table, bc)

            row_cells = table.add_row().cells
            row_cells[0].text = "Pressure"
            row_cells[1].text = bc.Pressure.UserString

            self.bcTurbulenceConfig(table, bc)

        elif bc.BoundarySubType == "volumetricFlowRateInlet":
            row_cells = table.add_row().cells
            row_cells[0].text = "Volume flow rate"
            row_cells[1].text = bc.VolFlowRate.UserString

            self.bcTurbulenceConfig(table, bc)

        elif bc.BoundarySubType == "massFlowRateInlet":
            row_cells = table.add_row().cells
            row_cells[0].text = "Mass flow rate"
            row_cells[1].text = bc.MassFlowRate.UserString

            self.bcTurbulenceConfig(table, bc)

        elif bc.BoundarySubType == "totalPressureInlet":
            row_cells = table.add_row().cells
            row_cells[0].text = "Pressure"
            row_cells[1].text = bc.Pressure.UserString

            self.bcTurbulenceConfig(table, bc)

    def bcOutletConfig(self, table, typeIndex, bc):
        if bc.BoundarySubType == "staticPressureOutlet":
            pass
        elif bc.BoundarySubType == "uniformVelocityOutlet":
            pass

    def bcOpenConfig(self, table, typeIndex, bc):
        if bc.BoundarySubType == "totalPressureOpening":
            pass
        elif bc.BoundarySubType == "farField":
            pass

    def bcConstraintConfig(self, typeIndex, table, bc):
        if bc.BoundarySubType == "cyclicAMI":
            pass

    def bcBaffleConfig(self, typeIndex, table, bc):
        pass

    def bcVelocityConfig(self, table, bc):
        if bc.VelocityIsCartesian:
            row_cells = table.add_row().cells
            row_cells[0].text = "Velocity"
            row_cells[1].text = "({}, {}, {})".format(
                bc.Ux.UserString,
                bc.Uy.UserString,
                bc.Uz.UserString,
            )
        else:
            row_cells = table.add_row().cells
            row_cells[0].text = "Velocity"
            row_cells[1].text = bc.VelocityMag.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Face"
            row_cells[1].text = bc.DirectionFace

            row_cells = table.add_row().cells
            row_cells[0].text = "Inward normal"
            row_cells[1].text = self.configState(bc.ReverseNormal)

    def bcTurbulenceConfig(self, table, bc):
        physicsModel = CfdTools.getPhysicsModel(self._analysis)
        turbulenceModel = (physicsModel.TurbulenceModel
                          if physicsModel.Turbulence == 'RANS' or physicsModel.Turbulence == 'DES'
                             or physicsModel.Turbulence == 'LES' else None)


        row_cells = table.add_row().cells
        row_cells[0].text = "Turbulence"
        # row_cells[1].text = TURBULENT_INLET_SPEC[turbulenceModel][0]

        if bc.TurbulenceInletSpecification in ['intensityAndLengthScale']:
            row_cells[1].text = "Intensity & Length scale"

            row_cells = table.add_row().cells
            row_cells[0].text = "Turbulence intensity (I)"
            row_cells[1].text = bc.TurbulenceIntensityPercentage.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Length scale (l)"
            row_cells[1].text = bc.TurbulenceLengthScale.UserString

        if bc.TurbulenceInletSpecification in ['TKEAndSpecDissipationRate']:
            row_cells[1].text = "Kinetic energy & Specific dissipation rate"

            row_cells = table.add_row().cells
            row_cells[0].text = "Turbulence kinetic energy (k)"
            row_cells[1].text = bc.TurbulentKineticEnergy.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Specific dissipation rate (ω)"
            row_cells[1].text = bc.SpecificDissipationRate.UserString

        if bc.TurbulenceInletSpecification in ['TurbulentViscosityAndK']:
            row_cells[1].text = "Kinetic Energy & Turbulent viscosity"

            row_cells = table.add_row().cells
            row_cells[0].text = "Turbulence kinetic energy (k)"
            row_cells[1].text = bc.kEqnTurbulentKineticEnergy.UserString

            row_cells = table.add_row().cells
            row_cells[0].text = "Turbulent viscosity (v)"
            row_cells[1].text = bc.kEqnTurbulentViscosity.UserString

        # if bc.TurbulenceInletSpecification in ['TKEAndDissipationRate']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "ε"
        #     row_cells[1].text = initModel.epsilon.UserString

        # if bc.TurbulenceInletSpecification in ['TKEAndSpecDissipationRate', 'intensityAndLengthScale', 'TKESpecDissipationRateGammaAndReThetat', 'TKEAndSpecDissipationRate']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "ω"
        #     row_cells[1].text = initModel.omega.UserString

        # if bc.TurbulenceInletSpecification in ['TransportedNuTilda']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "ṽ"
        #     row_cells[1].text = initModel.nuTilda.UserString

        # if bc.TurbulenceInletSpecification in ['TKESpecDissipationRateGammaAndReThetat']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "γ"
        #     row_cells[1].text = initModel.gammaInt.UserString

        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "Reθ"
        #     row_cells[1].text = initModel.ReThetat.UserString

        # if bc.TurbulenceInletSpecification in ['TurbulentViscosity', 'TurbulentViscosity', 'TurbulentViscosityAndK']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "v"
        #     row_cells[1].text = initModel.nut.UserString

        # if bc.TurbulenceInletSpecification in ['TKEAndSpecDissipationRate', 'intensityAndLengthScale', 'TKESpecDissipationRateGammaAndReThetat', 'TKEAndSpecDissipationRate', 'TKEAndDissipationRate', 'TurbulentViscosityAndK']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "k"
        #     row_cells[1].text = initModel.k.UserString

        # if bc.TurbulenceInletSpecification in ['TKEAndDissipationRate']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "ε"
        #     row_cells[1].text = initModel.epsilon.UserString

        # if bc.TurbulenceInletSpecification in ['TKEAndSpecDissipationRate', 'intensityAndLengthScale', 'TKESpecDissipationRateGammaAndReThetat', 'TKEAndSpecDissipationRate']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "ω"
        #     row_cells[1].text = initModel.omega.UserString

        # if bc.TurbulenceInletSpecification in ['TransportedNuTilda']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "ṽ"
        #     row_cells[1].text = initModel.nuTilda.UserString

        # if bc.TurbulenceInletSpecification in ['TKESpecDissipationRateGammaAndReThetat']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "γ"
        #     row_cells[1].text = initModel.gammaInt.UserString

        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "Reθ"
        #     row_cells[1].text = initModel.ReThetat.UserString

        # if bc.TurbulenceInletSpecification in ['TurbulentViscosity', 'TurbulentViscosity', 'TurbulentViscosityAndK']:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = "v"
        #     row_cells[1].text = initModel.nut.UserString

    def configState(self, condition):
        if condition:
            return "True"
        else:
            return "False"

    def generateRuntime(self):
        self._document.add_heading('Run status', level=1)

        self._document.add_paragraph("Total run time: " + self._totalRunTime)

        table = self._document.add_table(rows=1, cols=4, style='Table Grid')
        table.autofit = True
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Angle of Attack'
        hdr_cells[1].text = 'Run Time'
        hdr_cells[2].text = 'Status'
        hdr_cells[3].text = 'Description'
        for angle in self._analysis.AOAList:
            row_cells = table.add_row().cells
            try:
                row_cells[0].text = self._runStatus[str(angle)][0]
                row_cells[1].text = self._runStatus[str(angle)][1]
                row_cells[2].text = self._runStatus[str(angle)][2]
                row_cells[3].text = self._runStatus[str(angle)][3]
            except KeyError:
                row_cells[0].text = str(angle)
                row_cells[1].text = ""
                row_cells[2].text = "Unknown"
                row_cells[3].text = ""

        self._document.add_page_break()

    def generateCP(self):
        self._document.add_heading('Center of Pressure', level=1)
        p = self._document.add_paragraph("")
        p.add_run("NOTE: It is not possible to calculate the center of pressure at an "\
                                    "angle of attack of 0. It must be calculated at multiple points close "\
                                    "to 0 and inferred using l'Hôpital's rule.").bold = True
        self._document.add_paragraph("Center of Pressure at multiple angles of attack with the rocket rotated "\
                                     "around the center point.")

        p = self._document.add_paragraph("CP = X0 - My / (Fz * cos(AOA) + Fx * sin(AOA))", style="Block Quotation")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = self._document.add_paragraph("X0 = {} mm".format(int(self._x0)), style="Block Quotation")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        x_values = []
        y_values = []
        count = len(self._analysis.AOAList)
        table = self._document.add_table(rows=5, cols=count+1, style='Simple Grid Columns')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Angle of Attack'
        for i in range(count):
            hdr_cells[i + 1].text = str(self._analysis.AOAList[i])
        data_cells = table.columns[0].cells
        data_cells[1].text = "My"
        data_cells[2].text = "Fx"
        data_cells[3].text = "Fz"
        data_cells[4].text = "CP"
        for i in range(count):
            data_cells = table.columns[i+1].cells
            angle = str(self._analysis.AOAList[i])
            if self._runStatus[str(angle)][2] == "Success":
                My = self.getMomentY(angle) #float(self._moments[angle][-1][2])
                Fx = self.getForceX(angle) #float(self._forces[angle][-1][1])
                Fz = self.getForceZ(angle) #float(self._forces[angle][-1][3])
                aoa = math.radians(float(self._analysis.AOAList[i]))
                # X = X0 - My / (Fz * cos(AOA) + Fx * sin(AOA))
                x = (My / (Fz * math.cos(aoa) + Fx * math.sin(aoa)))
                x = self._x0 - (x * 1000.0) # Meters to mm

                x_values.append(float(self._analysis.AOAList[i]))
                y_values.append(x)
                data_cells[1].text = "{:#.3g}".format(My)
                data_cells[2].text = "{:#.3g}".format(Fx)
                data_cells[3].text = "{:#.3g}".format(Fz)
                data_cells[4].text = "{} mm".format(int(x))
            else:
                data_cells[1].text = ""
                data_cells[2].text = ""
                data_cells[3].text = ""
                data_cells[4].text = ""

        if count > 1:
            self._document.add_paragraph()
            self.cpGraph(x_values, y_values)

        self._document.add_page_break()

    def getMomentY(self, angle):
        return self.getAverage(self._moments[angle], 2)

    def getForceX(self, angle):
        return self.getAverage(self._forces[angle], 1)

    def getForceZ(self, angle):
        return self.getAverage(self._forces[angle], 3)

    def getCD(self, angle):
        return self.getAverage(self._coefficients[angle], 1)

    def getCL(self, angle):
        return self.getAverage(self._coefficients[angle], 4)

    def getAverage(self, list, index):
        if len(list) > self._analysis.AverageLastN:
            list = list[-self._analysis.AverageLastN:]
        average = 0.0
        for entry in list:
            average += float(entry[index])
        return (average / len(list))

    def cpGraph(self, x_values, y_values):
        graphPath = os.path.join(CfdTools.getOutputPath(self._analysis), 'cpGraph.png')

        # Turn interactive plotting off
        plt.ioff()
        plt.rcParams['figure.constrained_layout.use'] = True
        plt.rcParams["figure.facecolor"] = 'white'
        plt.rcParams["figure.edgecolor"] = 'white'
        plt.rcParams["axes.facecolor"] = 'white'

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, y_values, label="$C_P$")
        sub.set_xlabel("Angle of Attack (degrees)")
        sub.set_ylabel("Center of Pressure (mm)")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))

    def generateCD(self):
        self._document.add_heading('Lift and Drag', level=1)
        p = self._document.add_paragraph("Lift and Drag coefficients at multiple angles of attack with the rocket rotated "\
                                     "around the center point.")

        count = len(self._analysis.AOAList)
        table = self._document.add_table(rows=3, cols=count+1, style='Simple Grid Columns')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Angle of Attack'
        for i in range(count):
            hdr_cells[i + 1].text = str(self._analysis.AOAList[i])
        data_cells = table.columns[0].cells
        # data_cells[1].text = "Projected flow stream area"
        # data_cells[2].text = "Lift area"
        data_cells[1].text = "Cd"
        data_cells[2].text = "Cl"

        x_values = []
        cd_values = []
        cl_values = []
        for i in range(count):
            data_cells = table.columns[i+1].cells
            angle = str(self._analysis.AOAList[i])
            if self._runStatus[str(angle)][2] == "Success":
                aoa = float(self._analysis.AOAList[i])
                Cd = self.getCD(angle) #float(self._coefficients[angle][1])
                Cl = self.getCL(angle) #float(self._coefficients[angle][4])
                solid = applyTranslations(self._analysis.Shape, self._x0, aoa, self._rotation)
                area = calcFrontalArea(solid)

                x_values.append(float(self._analysis.AOAList[i]))
                cd_values.append(Cd)
                cl_values.append(Cl)
                # data_cells[1].text = FreeCAD.Units.Quantity("{} mm^2".format(area)).UserString
                # data_cells[2].text = ""
                data_cells[1].text = "{:#.3g}".format(Cd)
                data_cells[2].text = "{:#.3g}".format(Cl)
            else:
                data_cells[1].text = ""
                data_cells[2].text = ""
                # data_cells[3].text = ""
                # data_cells[4].text = ""

        if count > 1:
            self._document.add_paragraph()
            self.cdGraph(x_values, cd_values, cl_values)

        self._document.add_page_break()

    def cdGraph(self, x_values, cd_values, cl_values):
        graphPath = os.path.join(CfdTools.getOutputPath(self._analysis), 'cdGraph.png')

        # Turn interactive plotting off
        plt.ioff()
        plt.rcParams['figure.constrained_layout.use'] = True
        plt.rcParams["figure.facecolor"] = 'white'
        plt.rcParams["figure.edgecolor"] = 'white'
        plt.rcParams["axes.facecolor"] = 'white'

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, cd_values, label="$C_D$")
        sub.plot(x_values, cl_values, label="$C_L$")
        sub.set_xlabel("Angle of Attack (degrees)")
        sub.set_ylabel("Coefficient")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))

    def generateRunData(self):
        self._document.add_heading('Run Data', level=1)
        for angle in self._analysis.AOAList:
            if self._runStatus[str(angle)][2] == "Success" or self._runStatus[str(angle)][3] == "Solver error":
                self._document.add_heading('Angle of Attack={}'.format(str(angle)), level=1)
                self.generateResidualDataAt(angle)
                self.generateForceDataAt(angle)
                self.generateMomentDataAt(angle)
                self.generateCoefficientDataAt(angle)

                # Page break for all but the last entry
                if angle != self._analysis.AOAList[-1]:
                    self._document.add_page_break()

    def generateResidualDataAt(self, angle):
        graphPath = os.path.join(CfdTools.getOutputPath(self._analysis), 'residualGraph.png')

        # Turn interactive plotting off
        plt.ioff()
        plt.rcParams['figure.constrained_layout.use'] = True
        plt.rcParams["figure.facecolor"] = 'white'
        plt.rcParams["figure.edgecolor"] = 'white'
        plt.rcParams["axes.facecolor"] = 'white'

        residuals = self.readResidual(angle)

        try:
            x_values = residuals["time"]
            ux_values = residuals["UxResiduals"]
            uy_values = residuals["UyResiduals"]
            uz_values = residuals["UzResiduals"]
            p_values = residuals["pResiduals"]
            k_values = residuals["kResiduals"]
            omega_values = residuals["omegaResiduals"]
        except IndexError:
            # If the entry doesn't exist we're unable to draw the graph
            return

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, ux_values, label="$U_x$")
        sub.plot(x_values, uy_values, label="$U_y$")
        sub.plot(x_values, uz_values, label="$U_z$")
        sub.plot(x_values, p_values, label="$p$")
        sub.plot(x_values, k_values, label="$k$")
        sub.plot(x_values, omega_values, label="$\\omega$")
        sub.set_xlabel("Iteration")
        sub.set_ylabel("Residual")
        sub.set_yscale('log')
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))
        self._document.add_paragraph() # Spacer

    def readResidual(self, angle):
        # Avoid circular import
        from Rocket.cfd.Ui.TaskPanelMultiCFD import FoamRunner

        dirname = "case_aoa_{}".format(angle)
        path = os.path.join(CfdTools.getOutputPath(self._analysis), dirname,
                            "log.simpleFoam")

        runner = FoamRunner(self._analysis)
        with open(path, "r") as file:
            for line in file:
                runner.processOutput(line)

        residuals = {}
        residuals["time"] = runner.time
        residuals["UxResiduals"] = runner.UxResiduals
        residuals["UyResiduals"] = runner.UyResiduals
        residuals["UzResiduals"] = runner.UzResiduals
        residuals["pResiduals"] = runner.pResiduals
        residuals["kResiduals"] = runner.kResiduals
        residuals["omegaResiduals"] = runner.omegaResiduals

        return residuals

    def generateForceDataAt(self, angle):
        graphPath = os.path.join(CfdTools.getOutputPath(self._analysis), 'forceGraph.png')

        # Turn interactive plotting off
        plt.ioff()
        plt.rcParams['figure.constrained_layout.use'] = True
        plt.rcParams["figure.facecolor"] = 'white'
        plt.rcParams["figure.edgecolor"] = 'white'
        plt.rcParams["axes.facecolor"] = 'white'

        # Get the force entries
        x_values = self.getIntColumn(self._forces[str(angle)], 0)
        fx_values = self.getColumn(self._forces[str(angle)], 1)
        fy_values = self.getColumn(self._forces[str(angle)], 2)
        fz_values = self.getColumn(self._forces[str(angle)], 3)
        px_values = self.getColumn(self._forces[str(angle)], 4)
        py_values = self.getColumn(self._forces[str(angle)], 5)
        pz_values = self.getColumn(self._forces[str(angle)], 6)
        vx_values = self.getColumn(self._forces[str(angle)], 7)
        vy_values = self.getColumn(self._forces[str(angle)], 8)
        vz_values = self.getColumn(self._forces[str(angle)], 9)

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, fx_values, label="$F_X$")
        sub.plot(x_values, fy_values, label="$F_Y$")
        sub.plot(x_values, fz_values, label="$F_Z$")
        sub.set_xlabel("Iteration")
        sub.set_ylabel("Force [N]")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))
        self._document.add_paragraph() # Spacer

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, px_values, label="$F_X$ (pressure)")
        sub.plot(x_values, py_values, label="$F_Y$ (pressure)")
        sub.plot(x_values, pz_values, label="$F_Z$ (pressure)")
        sub.plot(x_values, vx_values, label="$F_X$ (viscous)")
        sub.plot(x_values, vy_values, label="$F_Y$ (viscous)")
        sub.plot(x_values, vz_values, label="$F_Z$ (viscous)")
        sub.set_xlabel("Iteration")
        sub.set_ylabel("Force [N]")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))
        self._document.add_paragraph() # Spacer

    def generateMomentDataAt(self, angle):
        graphPath = os.path.join(CfdTools.getOutputPath(self._analysis), 'momentGraph.png')

        # Turn interactive plotting off
        plt.ioff()
        plt.rcParams['figure.constrained_layout.use'] = True
        plt.rcParams["figure.facecolor"] = 'white'
        plt.rcParams["figure.edgecolor"] = 'white'
        plt.rcParams["axes.facecolor"] = 'white'

        # Get the force entries
        x_values = self.getIntColumn(self._moments[str(angle)], 0)
        fx_values = self.getColumn(self._moments[str(angle)], 1)
        fy_values = self.getColumn(self._moments[str(angle)], 2)
        fz_values = self.getColumn(self._moments[str(angle)], 3)
        px_values = self.getColumn(self._moments[str(angle)], 4)
        py_values = self.getColumn(self._moments[str(angle)], 5)
        pz_values = self.getColumn(self._moments[str(angle)], 6)
        vx_values = self.getColumn(self._moments[str(angle)], 7)
        vy_values = self.getColumn(self._moments[str(angle)], 8)
        vz_values = self.getColumn(self._moments[str(angle)], 9)

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, fx_values, label="$M_X$")
        sub.plot(x_values, fy_values, label="$M_Y$")
        sub.plot(x_values, fz_values, label="$M_Z$")
        sub.set_xlabel("Iteration")
        sub.set_ylabel("Moment")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))
        self._document.add_paragraph() # Spacer

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, px_values, label="$M_X$ (pressure)")
        sub.plot(x_values, py_values, label="$M_Y$ (pressure)")
        sub.plot(x_values, pz_values, label="$M_Z$ (pressure)")
        sub.plot(x_values, vx_values, label="$M_X$ (viscous)")
        sub.plot(x_values, vy_values, label="$M_Y$ (viscous)")
        sub.plot(x_values, vz_values, label="$M_Z$ (viscous)")
        sub.set_xlabel("Iteration")
        sub.set_ylabel("Moment")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))
        self._document.add_paragraph() # Spacer

    def generateCoefficientDataAt(self, angle):
        graphPath = os.path.join(CfdTools.getOutputPath(self._analysis), 'coefficientGraph.png')

        # Turn interactive plotting off
        plt.ioff()
        plt.rcParams['figure.constrained_layout.use'] = True
        plt.rcParams["figure.facecolor"] = 'white'
        plt.rcParams["figure.edgecolor"] = 'white'
        plt.rcParams["axes.facecolor"] = 'white'

        # Get the force entries
        x_values = self.getIntColumn(self._coefficients[str(angle)], 0)
        cd_values = self.getColumn(self._coefficients[str(angle)], 1)
        cl_values = self.getColumn(self._coefficients[str(angle)], 4)

        # Create a new figure, plot into it, then close it so it never gets displayed
        fig = plt.figure(figsize=(5,3))
        canvas = FigureCanvas(fig)
        sub = canvas.figure.subplots()
        sub.plot(x_values, cd_values, label="$C_D$")
        sub.plot(x_values, cl_values, label="$C_L$")
        sub.set_xlabel("Iteration")
        sub.set_ylabel("Coefficient")
        sub.grid(visible=True)
        sub.legend()

        plt.savefig(graphPath)
        plt.close(fig)

        self._document.add_picture(graphPath, width=Inches(6.0))
        # self._document.add_paragraph() # Spacer

    def getIntColumn(self, table, index):
        column = []
        for entry in table:
            column.append(int(entry[index]))
        return column

    def getColumn(self, table, index):
        column = []
        for entry in table:
            column.append(float(entry[index]))
        return column
